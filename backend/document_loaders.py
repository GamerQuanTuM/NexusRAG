import os
import tempfile
from typing import List, Optional, Dict, Any
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
    WebBaseLoader
)
from langchain_core.documents import Document
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Lightweight Word doc support
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("'python-docx' not installed. Word document support via lightweight loader is disabled.")

# Optional heavy loaders (require 'unstructured' package)
try:
    from langchain_community.document_loaders import UnstructuredMarkdownLoader
    from langchain_community.document_loaders import UnstructuredWordDocumentLoader
    from langchain_community.document_loaders import UnstructuredPowerPointLoader
    from langchain_community.document_loaders import UnstructuredExcelLoader
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False
    logger.warning("'unstructured' not installed. Word, PPT, Excel, and Markdown loaders are disabled.")


class DocumentLoader:
    """Handles loading documents from various file formats"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.txt': 'text',
        '.md': 'markdown',
        '.csv': 'csv',
        '.doc': 'word',
        '.docx': 'word',
        '.ppt': 'powerpoint',
        '.pptx': 'powerpoint',
        '.xls': 'excel',
        '.xlsx': 'excel',
    }
    
    @classmethod
    def get_file_type(cls, file_path: str) -> Optional[str]:
        """Get file type based on extension"""
        ext = Path(file_path).suffix.lower()
        return cls.SUPPORTED_EXTENSIONS.get(ext)
    
    @classmethod
    def load_document(cls, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Load a single document from file path"""
        file_type = cls.get_file_type(file_path)
        
        if not file_type:
            raise ValueError(f"Unsupported file type: {file_path}")
        
        logger.info(f"Loading {file_type} document: {file_path}")
        
        try:
            if file_type == 'pdf':
                loader = PyPDFLoader(file_path)
            elif file_type == 'text':
                loader = TextLoader(file_path, encoding='utf-8')
            elif file_type == 'markdown':
                if not HAS_UNSTRUCTURED:
                    # Fallback: read markdown as plain text
                    loader = TextLoader(file_path, encoding='utf-8')
                else:
                    loader = UnstructuredMarkdownLoader(file_path)
            elif file_type == 'csv':
                loader = CSVLoader(file_path)
            elif file_type == 'word':
                if HAS_UNSTRUCTURED:
                    loader = UnstructuredWordDocumentLoader(file_path)
                elif HAS_DOCX:
                    # Lightweight fallback using python-docx
                    doc = docx.Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    docs = [Document(page_content=text, metadata=metadata or {})]
                    if metadata:
                        for d in docs:
                            d.metadata.update(metadata)
                    return docs
                else:
                    raise ValueError("Word document support requires either 'unstructured' or 'python-docx'. Please upload a .pdf or .txt file instead.")
            elif file_type == 'powerpoint':
                if not HAS_UNSTRUCTURED:
                    raise ValueError("PowerPoint support requires the 'unstructured' package. Please upload a .pdf or .txt file instead.")
                loader = UnstructuredPowerPointLoader(file_path)
            elif file_type == 'excel':
                if not HAS_UNSTRUCTURED:
                    raise ValueError("Excel support requires the 'unstructured' package. Please upload a .pdf or .txt file instead.")
                loader = UnstructuredExcelLoader(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            documents = loader.load()
            
            # Add custom metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            logger.info(f"Successfully loaded {len(documents)} document chunks from {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            raise
    
    @classmethod
    def load_documents_from_directory(
        cls, 
        directory_path: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load all supported documents from a directory"""
        all_documents = []
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Directory not found: {directory_path}")
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and cls.get_file_type(str(file_path)):
                try:
                    documents = cls.load_document(str(file_path), metadata)
                    all_documents.extend(documents)
                except Exception as e:
                    logger.warning(f"Skipping file {file_path}: {e}")
        
        logger.info(f"Loaded {len(all_documents)} total document chunks from directory {directory_path}")
        return all_documents
    
    @classmethod
    def load_from_url(cls, url: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Load document from URL"""
        logger.info(f"Loading document from URL: {url}")
        
        try:
            loader = WebBaseLoader(url)
            documents = loader.load()
            
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            logger.info(f"Successfully loaded {len(documents)} document chunks from URL")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading from URL {url}: {e}")
            raise
    
    @classmethod
    def load_from_bytes(
        cls, 
        file_bytes: bytes, 
        filename: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load document from bytes (for uploaded files)"""
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_file_path = tmp_file.name
        
        try:
            documents = cls.load_document(tmp_file_path, metadata)
            return documents
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_file_path)
            except OSError:
                pass